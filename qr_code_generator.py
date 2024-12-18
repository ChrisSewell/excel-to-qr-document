import os
import qrcode
import xlsxwriter
from PIL import Image, ImageDraw, ImageFont
import io
import openpyxl
import tempfile
import shutil
from docx import Document
from docx.shared import Inches, Pt, Mm
import argparse
import math
from tqdm import tqdm
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from rich.console import Console
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn, TimeElapsedColumn
from rich import print as rprint
import platform
import concurrent.futures
from queue import Queue
from threading import Lock
from rich.table import Table
import signal

# Define PDF page dimensions
LETTER_WIDTH, LETTER_HEIGHT = letter

console = Console()

def signal_handler(sig, frame):
    console.print("\n\n[yellow]Process interrupted by user. Cleaning up...[/yellow]")
    exit(0)

signal.signal(signal.SIGINT, signal_handler)

def generate_single_qr(args):
    row_num, cell_value, qr_width, img_width, img_height, tmpdirname, args_verbose = args
    
    if args_verbose:
        console.print(f"\n[cyan]Processing row {row_num + 1}[/cyan]")
        console.print(f"[dim]Text:[/dim] '{cell_value}'")
        console.print(f"[dim]Text length:[/dim] {len(str(cell_value))}")
    
    try:
        # Generate QR code with larger dimensions
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4
        )
        qr.add_data(str(cell_value))
        qr.make(fit=True)
        qr_img = qr.make_image(fill_color="black", back_color="white")
        
        if args_verbose:
            console.print(f"[dim]Created QR code for:[/dim] {cell_value}")
        
        # Create larger base image
        base_width = 1000  # Fixed base width
        base_height = int(base_width * 1.2)  # Add 20% for text
        img_with_text = Image.new('RGB', (base_width, base_height), color='white')
        
        if args_verbose:
            console.print(f"[dim]Image dimensions:[/dim] {base_width}x{base_height}")
        
        # Add text
        text = str(cell_value)
        font_size = min(60, int(base_width / max(len(text), 10)))  # Scale font size based on text length
        text_height = int(base_height * 0.2)  # 20% of height for text
        
        # Font handling
        system_fonts = []
        if platform.system() == 'Windows':
            windows_font_paths = [
                "C:/Windows/Fonts/arial.ttf",
                "C:/Windows/Fonts/segoeui.ttf",
                "C:/Windows/Fonts/calibri.ttf",
                "C:/Windows/Fonts/tahoma.ttf",
            ]
            system_fonts.extend(windows_font_paths)
        elif platform.system() == 'Darwin':
            mac_font_paths = [
                "/System/Library/Fonts/Helvetica.ttc",
                "/Library/Fonts/Arial.ttf",
                "/System/Library/Fonts/SFNSText.ttf",
                "/System/Library/Fonts/Supplemental/Arial.ttf",
            ]
            system_fonts.extend(mac_font_paths)
        else:
            linux_font_paths = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/TTF/arial.ttf",
                "/usr/share/fonts/liberation/LiberationSans-Regular.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
                "/usr/share/fonts/ubuntu/Ubuntu-R.ttf",
            ]
            system_fonts.extend(linux_font_paths)
        
        font = None
        for font_path in system_fonts:
            try:
                if os.path.exists(font_path):
                    font = ImageFont.truetype(font_path, font_size)
                    break
            except Exception:
                continue
        
        if font:
            # Use system font
            draw = ImageDraw.Draw(img_with_text)
            left, top, right, bottom = font.getbbox(text)
            text_width = right - left
            text_position = ((base_width - text_width) // 2, text_height // 2)
            draw.text(text_position, text, fill='black', font=font)
        else:
            # Use default font with scaling
            default_font = ImageFont.load_default()
            
            # Create temporary image for text
            temp_img = Image.new('RGB', (base_width, text_height), color='white')
            temp_draw = ImageDraw.Draw(temp_img)
            
            # Draw text
            temp_draw.text((5, 5), text, font=default_font, fill='black')
            
            # Get text dimensions
            bbox = temp_draw.textbbox((0, 0), text, font=default_font)
            text_width = bbox[2] - bbox[0]
            text_actual_height = bbox[3] - bbox[1]
            
            # Scale text to fit
            scale_factor = min(
                (base_width * 0.8) / text_width,  # 80% of width
                (text_height * 0.8) / text_actual_height  # 80% of text area height
            )
            
            scaled_width = int(text_width * scale_factor)
            scaled_height = int(text_actual_height * scale_factor)
            
            # Resize text image
            text_img = temp_img.crop((0, 0, text_width + 10, text_actual_height + 10))
            text_img = text_img.resize(
                (scaled_width, scaled_height),
                Image.Resampling.LANCZOS
            )
            
            # Center the text
            text_x = (base_width - scaled_width) // 2
            text_y = (text_height - scaled_height) // 2
            img_with_text.paste(text_img, (text_x, text_y))
        
        # Add QR code
        qr_size = base_width
        qr_img = qr_img.resize((qr_size, qr_size))
        img_with_text.paste(qr_img, (0, text_height))
        
        # Save temporary file
        temp_file = os.path.join(tmpdirname, f'temp_qr_{row_num}.png')
        img_with_text.save(temp_file, quality=95)
        
        if args_verbose:
            console.print(f"[dim]Saved temporary file:[/dim] {temp_file}")
        
        return row_num, temp_file
        
    except Exception as e:
        if args_verbose:
            console.print(f"[red]Error processing row {row_num + 1}:[/red] {str(e)}")
        return row_num, None

def create_page(args):
    try:
        qr_files, page_num, hi_res_width, hi_res_height, hi_res_margin, hi_res_qr, wide, page_width, page_height = args
        
        # Create new page
        current_page = Image.new('RGB', (hi_res_width, hi_res_height), 'white')
        x = hi_res_margin
        y = hi_res_margin
        items_in_row = 0
        
        # Process QR codes for this page
        for temp_file in qr_files:
            if not os.path.exists(temp_file):
                continue
                
            # Load and resize QR code
            qr_img = Image.open(temp_file)
            qr_img = qr_img.resize((hi_res_qr, hi_res_qr), Image.Resampling.LANCZOS)
            
            # Paste QR code
            current_page.paste(qr_img, (x, y))
            
            # Update position
            items_in_row += 1
            if items_in_row >= wide:
                y += hi_res_qr
                x = hi_res_margin
                items_in_row = 0
            else:
                x += hi_res_qr
        
        # Scale down the page
        final_page = current_page.resize((page_width, page_height), Image.Resampling.LANCZOS)
        return page_num, final_page
        
    except Exception as e:
        console.print(f"[yellow]Warning:[/yellow] Failed to create page {page_num}: {str(e)}")
        return page_num, None

def generate_pdf(qr_files, wide, output_file, verbose=False):
    try:
        if verbose:
            console.print("\n[cyan]Initializing PDF generation...[/cyan]")
        
        # Initialize PDF with letter size
        width, height = LETTER_WIDTH, LETTER_HEIGHT
        c = canvas.Canvas(output_file, pagesize=(width, height))
        
        if verbose:
            console.print(f"[dim]Page size:[/dim] {width}x{height} points")
            console.print(f"[dim]QR codes per row:[/dim] {wide}")
        
        # Set margins (1 inch)
        margin = inch
        usable_width = width - (2 * margin)
        usable_height = height - (2 * margin)
        
        # Calculate QR code size
        qr_width = usable_width / wide
        qr_height = qr_width * 1.2  # Add space for text
        
        # Starting position (top-left with margin)
        x = margin
        y = height - margin - qr_height
        items_in_row = 0
        
        page_count = 1
        for temp_file in qr_files:
            try:
                # Add image
                c.drawImage(
                    temp_file,
                    x, y,
                    width=qr_width,
                    height=qr_height,
                    preserveAspectRatio=True
                )
                
                # Update position
                items_in_row += 1
                if items_in_row >= wide:
                    # Move to next row
                    y -= qr_height
                    x = margin
                    items_in_row = 0
                    
                    # Check if we need a new page
                    if y < margin:
                        if verbose:
                            console.print(f"[dim]Starting page {page_count + 1}[/dim]")
                        c.showPage()
                        y = height - margin - qr_height
                        x = margin
                        items_in_row = 0
                        page_count += 1
                else:
                    # Move to next position in row
                    x += qr_width
                
            except Exception as e:
                if verbose:
                    console.print(f"[yellow]Warning:[/yellow] Failed to add image {temp_file}: {str(e)}")
                continue
        
        if verbose:
            console.print(f"[dim]Total pages:[/dim] {page_count}")
            console.print("[cyan]Saving PDF...[/cyan]")
        
        c.save()
        return True
        
    except Exception as e:
        if verbose:
            console.print(f"[red]Error in PDF generation:[/red] {str(e)}")
        return False

def show_rich_help():
    console.print("\n[bold cyan]QR Code Generator for Excel Data[/bold cyan]")
    console.print("-------------------------------")
    console.print("This script generates QR codes from Excel data and arranges them in a document.\n")
    
    console.print("[bold cyan]Features:[/bold cyan]")
    console.print("• Interactive prompts for all options")
    console.print("• Supports both DOCX and PDF output")
    console.print("• Parallel processing for faster generation")
    console.print("• Preview data before column selection")
    console.print("• Configurable grid layout")
    console.print("• Progress tracking with elapsed time")
    console.print("• Detailed or minimal progress display\n")
    
    console.print("[bold green]Basic Usage:[/bold green]")
    console.print("    python qr_code_generator.py\n")
    
    console.print("[bold yellow]Command Line Options:[/bold yellow]")
    console.print("    -i, --input    Specify Excel file to use")
    console.print("    -c, --column   Select column number (1-based)")
    console.print("    -w, --wide     Number of QR codes per row (1-10)")
    console.print("    -t, --type     Output format (docx/pdf)")
    console.print("    --header       Skip header row prompt")
    console.print("    -v, --verbose  Show detailed progress\n")
    
    console.print("[bold cyan]Examples:[/bold cyan]")
    console.print("    # Interactive mode:")
    console.print("    python qr_code_generator.py\n")
    console.print("    # Specify all options:")
    console.print("    python qr_code_generator.py -i data.xlsx -c 2 -w 3 --header -t pdf\n")
    console.print("    # Common combinations:")
    console.print("    python qr_code_generator.py -i data.xlsx -c 2    # Use specific file and column")
    console.print("    python qr_code_generator.py -w 4 -t pdf          # 4-wide grid in PDF format")
    console.print("    python qr_code_generator.py --header -v          # Skip header, show details\n")
    
    console.print("[bold yellow]Notes:[/bold yellow]")
    console.print("• Output is saved to the 'output' directory")
    console.print("• Temporary files are automatically cleaned up")
    console.print("• Use Ctrl+C to cancel at any time")

# Add this class after imports
class RichHelpFormatter(argparse.HelpFormatter):
    def __init__(self, prog):
        super().__init__(prog)
        
    def format_help(self):
        console = Console()
        show_rich_help()
        return ""  # Return empty string since we've already printed the help

# Update the parser creation
parser = argparse.ArgumentParser(
    formatter_class=RichHelpFormatter,
    add_help=False  # Disable default help
)

# Add custom help arguments
parser.add_argument(
    '-h', '--help',
    action='help',
    help=argparse.SUPPRESS  # Hide from help output since we handle it in the formatter
)

# Remove the --help-rich argument since it's no longer needed
# parser.add_argument('--help-rich'...) can be removed

# Remove the help-rich handling code
# args, remaining = parser.parse_known_args()
# if args.help_rich:
#     show_rich_help()
#     exit(0)

parser.add_argument(
    '-w', '--wide',
    type=int,
    choices=range(1, 11),
    metavar='WIDTH',
    help='Number of QR codes per row (1-10, default: prompt user)'
)

parser.add_argument(
    '-t', '--type',
    choices=['docx', 'pdf'],
    default=None,
    help='Output file type (docx or pdf, default: prompt user)'
)

parser.add_argument(
    '-v', '--verbose',
    action='store_true',
    help='Show detailed progress instead of progress bar'
)

parser.add_argument(
    '-i', '--input',
    type=str,
    help='Input Excel file path'
)

parser.add_argument(
    '-c', '--column',
    type=int,
    help='Column number to use for QR codes (1-based)'
)

parser.add_argument(
    '--header',
    action='store_true',
    help='Specify if Excel file has a header row'
)

args = parser.parse_args()

try:
    console.print("\n[bold cyan]QR Code Generator[/bold cyan]")
    console.print("[bold cyan]═[/bold cyan]" * 50)
    console.print("\n[yellow]This tool will help you generate QR codes from your Excel data.[/yellow]")
    
    def clear_terminal():
        """Clear the terminal screen based on the operating system."""
        if platform.system().lower() == "windows":
            os.system('cls')
        else:
            os.system('clear')

    clear_terminal()

    # Handle file selection
    if args.input:
        if not os.path.exists(args.input) or not args.input.endswith('.xlsx'):
            console.print(f"[bold red]Error:[/bold red] Invalid Excel file: {args.input}")
            exit()
        selected_file = args.input
        console.print(f"[green]Using Excel file:[/green] {selected_file}")
    else:
        # Existing file selection code
        xlsx_files = [f for f in os.listdir() if f.endswith('.xlsx') and not f.startswith('~$')]
        if len(xlsx_files) == 0:
            console.print("\n[bold red]Error:[/bold red] No Excel files found in the current directory.")
            exit()
        elif len(xlsx_files) == 1:
            selected_file = xlsx_files[0]
            console.print(f"\n[green]Found Excel file:[/green] {selected_file}")
        else:
            console.print("\n[yellow]Multiple Excel files found:[/yellow]")
            for i, file in enumerate(xlsx_files, 1):
                console.print(f"  {i}. {file}")
            while True:
                try:
                    choice = int(console.input("\nEnter the number of your choice (1-{len(xlsx_files)}): ")) - 1
                    if 0 <= choice < len(xlsx_files):
                        selected_file = xlsx_files[choice]
                        console.print(f"\n[green]Selected:[/green] {selected_file}")
                        break
                    else:
                        console.print(f"[red]Please enter a number between 1 and {len(xlsx_files)}[/red]")
                except ValueError:
                    console.print("[red]Please enter a valid number[/red]")

    # Handle header flag
    has_header = args.header
    if not args.header and not args.column:  # Only prompt if column not specified
        while True:
            header_choice = console.input("\n[yellow]Does this Excel file have a header row? ([green]y[/green]/[red]n[/red]): [/yellow]").lower()
            if header_choice in ['y', 'n']:
                has_header = header_choice == 'y'
                break
            else:
                console.print("[red]Please enter 'y' for yes or 'n' for no[/red]")
    start_row = 2 if has_header else 1

    # Load workbook and get worksheet info
    src_wb = openpyxl.load_workbook(selected_file, read_only=True, data_only=True)
    src_ws = src_wb.active
    max_col = src_ws.max_column

    # Handle column selection
    if args.column:
        if 1 <= args.column <= max_col:
            selected_column = args.column
            if has_header:
                header_value = src_ws.cell(row=1, column=selected_column).value
                console.print(f"[green]Using column {selected_column} ({header_value})[/green]")
            else:
                console.print(f"[green]Using column {selected_column}[/green]")
        else:
            console.print(f"[bold red]Error:[/bold red] Column number must be between 1 and {max_col}")
            exit()
    else:
        # Get the maximum column with data
        max_preview_rows = min(6, src_ws.max_row)  # Show first 5 rows (plus header if exists)
        
        # Create a list to store the preview data
        preview_data = []
        header_row = []
        
        # Get column letters for display
        column_letters = [openpyxl.utils.get_column_letter(i) for i in range(1, max_col + 1)]
        
        # Collect preview data
        for row in range(1, max_preview_rows + 1):
            row_data = []
            for col in range(1, max_col + 1):
                cell_value = src_ws.cell(row=row, column=col).value
                if cell_value is None:
                    cell_value = ""
                row_data.append(str(cell_value))
            if row == 1 and has_header:
                header_row = row_data
            else:
                preview_data.append(row_data)
        
        # Create and display the preview table
        preview_table = Table(title="\nFirst 5 rows of data")
        
        # Add column headers
        for idx, letter in enumerate(column_letters, 1):
            if has_header:
                column_name = f"{idx}. {letter} - {header_row[idx-1]}"
            else:
                column_name = f"{idx}. Column {letter}"
            preview_table.add_column(column_name)
        
        # Add data rows
        for row in preview_data:
            preview_table.add_row(*row)
        
        console.print(preview_table)
        
        # Get column selection
        while True:
            try:
                column_choice = int(console.input("\n[yellow]Enter the number of the column to use for QR codes:[/yellow] "))
                if 1 <= column_choice <= max_col:
                    selected_column = column_choice
                    break
                else:
                    console.print(f"[red]Please enter a number between 1 and {max_col}[/red]")
            except ValueError:
                console.print("[red]Please enter a valid number[/red]")

    # Collect valid rows
    with console.status("[bold yellow]Analyzing Excel data...[/bold yellow]"):
        valid_rows = []
        for row in range(start_row, src_ws.max_row + 1):
            cell_value = src_ws.cell(row=row, column=selected_column).value
            if cell_value:
                valid_rows.append((row - start_row, cell_value))
        
        print(f"Found {len(valid_rows)} valid rows in selected column")
        if len(valid_rows) == 0:
            console.print("[bold red]Error:[/bold red] No data found in selected column")
            exit()

    # Handle width selection
    if args.wide is None:
        while True:
            try:
                width_choice = int(console.input("\n[yellow]How many QR codes per row? (1-10):[/yellow] "))
                if 1 <= width_choice <= 10:
                    args.wide = width_choice
                    break
                else:
                    console.print("[red]Please enter a number between 1 and 10[/red]")
            except ValueError:
                console.print("[red]Please enter a valid number[/red]")

    # Handle output format selection
    if args.type is None:
        console.print("\n[yellow]Select output format:[/yellow]")
        console.print("  [green]1.[/green] Word Document (DOCX)")
        console.print("  [red]2.[/red] PDF Document")
        while True:
            try:
                format_choice = int(console.input("\nEnter number (1-2): "))
                if format_choice == 1:
                    args.type = 'docx'
                    break
                elif format_choice == 2:
                    args.type = 'pdf'
                    break
                else:
                    console.print("[red]Please enter either 1 or 2[/red]")
            except ValueError:
                console.print("[red]Please enter a valid number[/red]")

    # Define output file path before summary
    output_dir = 'output'
    output_file = os.path.join(output_dir, f"qr_codes.{args.type}")

    # Show configuration summary
    console.print("\n[bold cyan]Configuration Summary:[/bold cyan]")
    console.print(f"[green]• Input file:[/green] {selected_file}")
    console.print(f"[green]• Header row:[/green] {'Yes' if has_header else 'No'}")
    if has_header:
        header_value = src_ws.cell(row=1, column=selected_column).value or ''  # Use empty string if None
        header_text = f" ({header_value})" if header_value else ""
        console.print(f"[green]• Selected column:[/green] {selected_column}{header_text}")
    else:
        console.print(f"[green]• Selected column:[/green] {selected_column}")
    console.print(f"[green]• QR codes per row:[/green] {args.wide}")
    console.print(f"[green]• Output format:[/green] {args.type.upper()}")
    console.print(f"[green]• Output file:[/green] {output_file}")

    if not args.verbose:
        console.print("\n[dim]Processing will show a progress bar. Use -v for detailed output.[/dim]")

    # Create output directory if it doesn't exist
    with console.status("[bold yellow]Setting up output directory...[/bold yellow]"):
        os.makedirs(output_dir, exist_ok=True)
        
        # Clear the output directory
        for filename in os.listdir(output_dir):
            file_path = os.path.join(output_dir, filename)
            while True:
                try:
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                    break  # Break the loop if deletion was successful
                except Exception as e:
                    console.print(f"\n[bold red]Error:[/bold red] Cannot delete {file_path}")
                    console.print("[yellow]This is likely because the file is open in another program.[/yellow]")
                    console.print("Please:")
                    console.print("1. Close any programs that might have the file open")
                    console.print("2. Press Enter to try again")
                    console.print("3. Press Ctrl+C to cancel the operation")
                    try:
                        console.input("\n[yellow]Press Enter when ready to continue...[/yellow]")
                    except KeyboardInterrupt:
                        console.print("\n[yellow]Operation cancelled by user[/yellow]")
                        exit(0)

    # Create a temporary directory
    with tempfile.TemporaryDirectory() as tmpdirname:
        print(f"Created temporary directory: {tmpdirname}")
        
        try:
            # Create document and calculate dimensions
            if args.type == 'docx':
                doc = Document()
                section = doc.sections[0]
                
                # Set margins (0.5 inch)
                margin_mm = 12.7  # 0.5 inch in mm
                section.left_margin = Mm(margin_mm)
                section.right_margin = Mm(margin_mm)
                section.top_margin = Mm(margin_mm)
                section.bottom_margin = Mm(margin_mm)
                
                # Calculate QR code size
                page_width_mm = section.page_width.mm - (2 * margin_mm)
                qr_width_mm = page_width_mm / args.wide
                qr_width_inches = qr_width_mm / 25.4  # Convert mm to inches
                
                # Set up progress display
                if args.verbose:
                    console.print("\n[bold cyan]Processing Details:[/bold cyan]")
                else:
                    progress = Progress(
                        SpinnerColumn(),
                        TextColumn("[progress.description]{task.description}"),
                        BarColumn(complete_style="green"),
                        TaskProgressColumn(),
                        TextColumn("•"),
                        TimeElapsedColumn(),
                        console=console,
                    )
                    task = progress.add_task("[cyan]Generating QR codes...", total=len(valid_rows))
                
                # Process QR codes
                if not args.verbose:
                    progress.start()
                
                try:
                    # Generate QR codes in parallel
                    process_args = [
                        (row_num, cell_value, qr_width_mm, 1000, 1000, tmpdirname, args.verbose)
                        for row_num, cell_value in valid_rows
                    ]
                    
                    results = {}
                    with concurrent.futures.ThreadPoolExecutor(max_workers=min(8, len(valid_rows))) as executor:
                        future_to_row = {
                            executor.submit(generate_single_qr, args): args[0] 
                            for args in process_args
                        }
                        
                        for future in concurrent.futures.as_completed(future_to_row):
                            row_num = future_to_row[future]
                            try:
                                row_num, temp_file = future.result()
                                if temp_file:
                                    results[row_num] = temp_file
                                    if not args.verbose:
                                        progress.advance(task)
                            except Exception as e:
                                console.print(f"\n[red]Error processing row {row_num + 1}:[/red] {str(e)}")
                    
                    # Sort results and create document
                    all_temp_files = [results[row_num] for row_num in sorted(results.keys())]
                    current_row = []
                    for temp_file in all_temp_files:
                        current_row.append(temp_file)
                        if len(current_row) == args.wide:
                            table = doc.add_table(rows=1, cols=args.wide)
                            for idx, tf in enumerate(current_row):
                                cell = table.cell(0, idx)
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(tf, width=Inches(qr_width_inches))
                            current_row = []
                            doc.add_paragraph()
                    
                    # Save the document
                    doc.save(output_file)
                    console.print(f"\n[bold green]Success![/bold green] Document saved as: [blue]{output_file}[/blue]")
                    
                except Exception as e:
                    if not args.verbose:
                        progress.stop()
                    console.print(f"\n[bold red]Error:[/bold red] {str(e)}")
                    raise
                
            else:
                # PDF Generation
                if not args.verbose:
                    progress = Progress(
                        SpinnerColumn(),
                        TextColumn("[progress.description]{task.description}"),
                        BarColumn(complete_style="green"),
                        TaskProgressColumn(),
                        TextColumn("•"),
                        TimeElapsedColumn(),
                        console=console,
                    )
                    task = progress.add_task("[cyan]Generating QR codes...", total=len(valid_rows))
                    progress.start()  # Start the progress bar
                
                try:
                    # First generate all QR codes in parallel
                    process_args = [
                        (row_num, cell_value, 1000, 1000, 1200, tmpdirname, args.verbose)
                        for row_num, cell_value in valid_rows
                    ]
                    
                    results = {}
                    with concurrent.futures.ThreadPoolExecutor(max_workers=min(8, len(valid_rows))) as executor:
                        future_to_row = {
                            executor.submit(generate_single_qr, args): args[0] 
                            for args in process_args
                        }
                        
                        for future in concurrent.futures.as_completed(future_to_row):
                            row_num = future_to_row[future]
                            try:
                                row_num, temp_file = future.result()
                                if temp_file:
                                    results[row_num] = temp_file
                                    if not args.verbose:
                                        progress.advance(task)
                            except Exception as e:
                                console.print(f"\n[red]Error processing row {row_num + 1}:[/red] {str(e)}")
                    
                    # Sort results by row number
                    all_temp_files = [results[row_num] for row_num in sorted(results.keys())]
                    
                    # Generate PDF
                    if not args.verbose:
                        progress.update(task, description="[cyan]Creating PDF...")
                    
                    if generate_pdf(all_temp_files, args.wide, output_file, args.verbose):
                        if not args.verbose:
                            progress.stop()
                        console.print(f"\n[bold green]Success![/bold green] Document saved as: [blue]{output_file}[/blue]")
                    else:
                        if not args.verbose:
                            progress.stop()
                        console.print("\n[bold red]Error:[/bold red] Failed to generate PDF")
                        
                except Exception as e:
                    if not args.verbose:
                        progress.stop()
                    console.print(f"\n[bold red]Error:[/bold red] {str(e)}")
                    raise
                
        except Exception as e:
            console.print(f"\n[bold red]Error:[/bold red] {str(e)}")

except KeyboardInterrupt:
    console.print("\n\n[yellow]Process interrupted by user. Cleaning up...[/yellow]")
    exit(0)
except Exception as e:
    console.print(f"\n[bold red]Error:[/bold red] {str(e)}")
    exit(1)
finally:
    console.print("\n[bold green]Script completed[/bold green]")
